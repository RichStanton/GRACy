#!/usr/bin/env python3
"""Convert GRACy planning markdown files to professional Word documents."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
BRAND_DARK   = RGBColor(0x1A, 0x2E, 0x4A)   # dark navy  – headings / rules
BRAND_MID    = RGBColor(0x2E, 0x6D, 0xA4)   # steel blue – H2 accent bar
BRAND_LIGHT  = RGBColor(0xE8, 0xF2, 0xFA)   # pale blue  – table header bg
CODE_BG      = RGBColor(0xF4, 0xF4, 0xF4)   # light grey – code block bg
CODE_FG      = RGBColor(0x1E, 0x1E, 0x1E)   # near-black – code text
META_FG      = RGBColor(0x55, 0x55, 0x55)   # mid grey   – metadata lines
TABLE_BORDER = RGBColor(0xBB, 0xCC, 0xDD)   # blue-grey  – table borders
ALT_ROW      = RGBColor(0xF7, 0xFA, 0xFD)   # very pale  – alternating row
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

BODY_FONT    = "Calibri"
MONO_FONT    = "Courier New"


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _set_cell_bg(cell, colour: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_col = f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_col)
    tcPr.append(shd)


def _set_cell_border(cell, **borders):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if side in borders:
            b = OxmlElement(f"w:{side}")
            for k, v in borders[side].items():
                b.set(qn(k), v)
            tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_paragraph_border_bottom(para, colour: RGBColor):
    """Add a bottom border to a paragraph (used for H2 rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    hex_col = f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}"
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_col)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_para_shading(para, colour: RGBColor):
    """Add background shading to a paragraph (code blocks)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_col = f"{colour[0]:02X}{colour[1]:02X}{colour[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_col)
    pPr.append(shd)


def _set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _set_para_indent(para, left_twips=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    pPr.append(ind)


def _add_horizontal_rule(doc):
    para = doc.add_paragraph()
    _set_para_spacing(para, before=60, after=60)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{BRAND_MID[0]:02X}{BRAND_MID[1]:02X}{BRAND_MID[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


# ── Inline markdown parser ────────────────────────────────────────────────────

def _apply_inline(run_parent, text: str, base_size: Pt, base_colour: RGBColor,
                  bold=False, mono=False):
    """
    Parse **bold**, *italic*, `code`, and plain text segments in `text`,
    appending runs to `run_parent` (a paragraph).
    """
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before this match
        if m.start() > pos:
            _add_run(run_parent, text[pos:m.start()], base_size, base_colour,
                     bold=bold, mono=mono)
        if m.group(2):  # **bold**
            _add_run(run_parent, m.group(2), base_size, base_colour,
                     bold=True, mono=mono)
        elif m.group(3):  # *italic*
            _add_run(run_parent, m.group(3), base_size, base_colour,
                     italic=True, mono=mono)
        elif m.group(4):  # `code`
            _add_run(run_parent, m.group(4), base_size, RGBColor(0xC7, 0x25, 0x4E),
                     mono=True)
        pos = m.end()
    if pos < len(text):
        _add_run(run_parent, text[pos:], base_size, base_colour,
                 bold=bold, mono=mono)


def _add_run(para, text, size: Pt, colour: RGBColor,
             bold=False, italic=False, mono=False):
    run = para.add_run(text)
    run.font.name = MONO_FONT if mono else BODY_FONT
    run.font.size = size
    run.font.color.rgb = colour
    run.font.bold = bold
    run.font.italic = italic
    return run


# ── Document-level builders ───────────────────────────────────────────────────

def _setup_doc(doc):
    """Apply global document styles: margins, default paragraph font."""
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10)


def _add_title_block(doc, title: str, meta_lines: list[str]):
    """Render the document title and metadata block at the top."""
    # Title
    p = doc.add_paragraph()
    _set_para_spacing(p, before=0, after=80)
    run = p.add_run(title)
    run.font.name = BODY_FONT
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK

    # Thin accent rule under title
    _add_horizontal_rule(doc)

    # Metadata lines (bold key, normal value)
    for line in meta_lines:
        p = doc.add_paragraph()
        _set_para_spacing(p, before=20, after=20)
        if ":" in line:
            key, _, val = line.partition(":")
            r = p.add_run(key.lstrip("* ").strip() + ": ")
            r.font.name = BODY_FONT
            r.font.size = Pt(9)
            r.font.bold = True
            r.font.color.rgb = META_FG
            _apply_inline(p, val.strip(), Pt(9), META_FG)
        else:
            _apply_inline(p, line.strip("* "), Pt(9), META_FG)

    doc.add_paragraph()  # spacer


def _add_h1(doc, text):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=240, after=80)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    _add_paragraph_border_bottom(p, BRAND_MID)
    return p


def _add_h2(doc, text):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=200, after=60)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BRAND_MID
    return p


def _add_h3(doc, text):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=140, after=40)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BRAND_DARK
    return p


def _add_h4(doc, text):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=100, after=30)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = BRAND_DARK
    return p


def _add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    _set_para_spacing(p, before=30, after=30)
    if indent:
        _set_para_indent(p, indent * 360)
    _apply_inline(p, text, Pt(10), RGBColor(0x22, 0x22, 0x22))
    return p


def _add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    _set_para_spacing(p, before=20, after=20)
    _set_para_indent(p, 360 + level * 360)
    _apply_inline(p, text, Pt(10), RGBColor(0x22, 0x22, 0x22))
    return p


def _add_code_block(doc, lines: list[str]):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        _add_para_shading(p, CODE_BG)
        _set_para_indent(p, 360)
        before = 80 if i == 0 else 0
        after  = 80 if i == len(lines) - 1 else 0
        _set_para_spacing(p, before=before, after=after)
        run = p.add_run(line if line else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = CODE_FG


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        _set_cell_bg(cell, BRAND_DARK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        _set_para_spacing(p, before=40, after=40)
        _apply_inline(p, h.strip(), Pt(9), WHITE, bold=True)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = ALT_ROW if r_idx % 2 == 1 else WHITE
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            _set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            _set_para_spacing(p, before=30, after=30)
            _apply_inline(p, cell_text.strip(), Pt(9), RGBColor(0x22, 0x22, 0x22))

    # Column widths: distribute proportionally, wider first column
    total_width = Inches(6.0)
    if col_count == 1:
        col_widths = [total_width]
    elif col_count == 2:
        col_widths = [Inches(2.2), Inches(3.8)]
    elif col_count == 3:
        col_widths = [Inches(1.6), Inches(2.4), Inches(2.0)]
    elif col_count == 4:
        col_widths = [Inches(0.8), Inches(2.4), Inches(1.6), Inches(1.2)]
    else:
        w = total_width / col_count
        col_widths = [w] * col_count

    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[i] if i < len(col_widths) else Inches(1.0)

    doc.add_paragraph()  # spacer after table


# ── Main parser / state machine ───────────────────────────────────────────────

def _parse_table_line(line):
    """Return list of cell strings from a markdown table row, or None."""
    line = line.strip()
    if not line.startswith("|"):
        return None
    cells = [c.strip() for c in line.strip("|").split("|")]
    return cells


def _is_separator_row(cells):
    return all(re.match(r"^:?-+:?$", c) for c in cells if c)


def convert_md_to_docx(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    _setup_doc(doc)

    # ── Extract title and metadata from the top of the document ──────────────
    title = md_path.stem.replace("_", " ").title()
    meta_lines = []
    content_start = 0

    for i, line in enumerate(lines):
        stripped = line.strip()
        if i == 0 and stripped.startswith("# "):
            title = stripped[2:].strip()
            content_start = 1
        elif stripped.startswith("**") and ":" in stripped and i < 12:
            meta_lines.append(stripped.strip("*").strip())
            content_start = i + 1
        elif stripped == "---" and i < 12:
            content_start = i + 1
            break
        elif stripped == "" and i < 12:
            content_start = i + 1
        elif not stripped.startswith("**") and stripped and i > 0 and i < 12:
            break

    _add_title_block(doc, title, meta_lines)

    # ── State machine over remaining lines ────────────────────────────────────
    i = content_start
    in_code = False
    code_lines = []
    table_headers = None
    table_rows = []
    pending_table_line = None

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # ── Code block ────────────────────────────────────────────────────────
        if line.startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^-{3,}$", line):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("#### "):
            _add_h4(doc, line[5:])
            i += 1
            continue
        if line.startswith("### "):
            _add_h3(doc, line[4:])
            i += 1
            continue
        if line.startswith("## "):
            _add_h2(doc, line[3:])
            i += 1
            continue
        if line.startswith("# "):
            _add_h1(doc, line[2:])
            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────────────────
        cells = _parse_table_line(line) if line.startswith("|") else None

        if cells is not None:
            if table_headers is None:
                table_headers = cells
            elif _is_separator_row(cells):
                pass  # skip separator
            else:
                table_rows.append(cells)
            i += 1
            # peek: if next line is not a table row, flush
            next_line = lines[i].strip() if i < len(lines) else ""
            if not next_line.startswith("|"):
                if table_headers:
                    _add_table(doc, table_headers, table_rows)
                table_headers = None
                table_rows = []
            continue

        # ── Bullet lists ──────────────────────────────────────────────────────
        m = re.match(r"^(\s*)([-*]|\d+\.) (.+)$", raw)
        if m:
            indent = len(m.group(1)) // 2
            _add_bullet(doc, m.group(3), level=indent)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line:
            i += 1
            continue

        # ── Body paragraph ────────────────────────────────────────────────────
        _add_body(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  Saved: {out_path.name}")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    planning_dir = Path(__file__).parent / "planning"
    output_dir   = Path(__file__).parent / "output"

    skip = {"decisions.md"}
    targets = sorted(p for p in planning_dir.glob("*.md") if p.name not in skip)

    if not targets:
        print("No markdown files found.")
        sys.exit(1)

    print(f"Converting {len(targets)} file(s) to Word documents in {output_dir}/")
    for md_path in targets:
        stem = md_path.stem.replace("_", " ").title().replace(" ", "_")
        out_path = output_dir / f"{stem}.docx"
        convert_md_to_docx(md_path, out_path)

    print("Done.")


if __name__ == "__main__":
    main()
